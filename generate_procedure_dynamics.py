"""Génère le document Word de la procédure de connexion Dynamics 365."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# -- Styles de base --
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    h = doc.styles[f"Heading {level}"]
    h.font.name = "Calibri"
    h.font.color.rgb = RGBColor(0x1A, 0x56, 0x8E)

doc.styles["Heading 1"].font.size = Pt(18)
doc.styles["Heading 2"].font.size = Pt(14)
doc.styles["Heading 3"].font.size = Pt(12)


def add_paragraph(text, bold=False, italic=False, style_name=None):
    p = doc.add_paragraph(style=style_name)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def add_numbered_step(number, text, sub_items=None):
    p = doc.add_paragraph()
    run = p.add_run(f"{number}. ")
    run.bold = True
    p.add_run(text)
    if sub_items:
        for item in sub_items:
            sub = doc.add_paragraph(style="List Bullet 2")
            sub.add_run(item)
    return p


def add_info_box(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    pPr = p._element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    left = pBdr.makeelement(
        qn("w:left"),
        {
            qn("w:val"): "single",
            qn("w:sz"): "12",
            qn("w:space"): "8",
            qn("w:color"): "1A568E",
        },
    )
    pBdr.append(left)
    pPr.append(pBdr)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.italic = True
    return p


def add_simple_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            table.rows[r_idx + 1].cells[c_idx].text = val
    doc.add_paragraph()


# ============================================================
# CONTENU
# ============================================================

doc.add_heading("Connexion du site FNECS avec Dynamics 365", level=1)
p = doc.add_paragraph()
p.add_run("Actions à réaliser côté FNECS").bold = True

doc.add_paragraph()

add_paragraph(
    "Bonjour Jean-René,", italic=True
)
add_paragraph(
    "Dans le cadre de la refonte du site, nous avons besoin de connecter le nouveau site "
    "à votre Dynamics 365 pour synchroniser les données adhérents. Voici les 3 actions "
    "que nous avons besoin que vous réalisiez de votre côté depuis votre compte administrateur.",
    italic=True,
)
add_paragraph(
    "On vous propose de faire ça ensemble en visio (30 min environ) — ce document vous "
    "sert de résumé avant et après.",
    italic=True,
)

# -- ACTION 1 --
doc.add_heading("Action 1 — Créer une « carte d'identité » pour le site dans Azure", level=2)

add_paragraph(
    "C'est ce qui permettra au site de communiquer avec Dynamics de manière sécurisée et automatique."
)

add_numbered_step(1, "Allez sur portal.azure.com et connectez-vous avec votre compte admin FNECS")
add_numbered_step(
    2,
    'Dans la barre de recherche en haut, tapez « Inscriptions d\'applications » et cliquez sur le résultat',
)
add_numbered_step(3, "Cliquez sur + Nouvelle inscription")
add_numbered_step(
    4,
    "Remplissez :",
    [
        "Nom : Site FNECS",
        "Types de comptes : choisissez la première option (cet annuaire uniquement)",
        "Le reste : ne touchez à rien",
    ],
)
add_numbered_step(5, "Cliquez Inscrire")
add_numbered_step(
    6,
    "Vous arrivez sur une page avec deux codes — copiez-les et envoyez-les-nous :",
    [
        "« ID d'application (client) »",
        "« ID de l'annuaire (locataire) »",
    ],
)
add_numbered_step(7, "Dans le menu à gauche, cliquez sur Certificats et secrets")
add_numbered_step(
    8,
    "Cliquez + Nouveau secret client > description : Site FNECS > cliquez Ajouter",
)
add_numbered_step(
    9,
    "Copiez immédiatement la valeur qui apparaît (la colonne « Valeur », pas « ID secret ») "
    "— elle disparaît définitivement si vous quittez la page",
)

add_info_box(
    "Envoyez-nous ces 3 informations (les 2 identifiants + le secret) par un canal sécurisé "
    "— pas par email en clair. Un message privé Teams ou un lien de partage sécurisé conviennent."
)

# -- ACTION 2 --
doc.add_heading("Action 2 — Autoriser le site à accéder aux données Dynamics", level=2)

add_numbered_step(1, "Allez sur admin.powerplatform.microsoft.com")
add_numbered_step(2, "Dans le menu à gauche, cliquez sur Environnements")

add_info_box(
    "Si vous ne voyez pas le menu « Environnements », c'est probablement un problème de droits. "
    "Contactez-nous et on réglera ça ensemble en visio."
)

add_numbered_step(
    3, "Cliquez sur votre environnement (celui où se trouvent vos adhérents)"
)
add_numbered_step(4, "Cliquez sur Paramètres en haut de la page")
add_numbered_step(
    5, "Cliquez sur Utilisateurs + autorisations puis sur Utilisateurs d'application"
)
add_numbered_step(6, "Cliquez sur + Nouvel utilisateur d'application")
add_numbered_step(
    7,
    "Cliquez sur + Ajouter une application > cherchez Site FNECS > sélectionnez-le > cliquez Ajouter",
)
add_numbered_step(8, "Choisissez votre division")
add_numbered_step(
    9,
    "Pour les rôles de sécurité : on choisira ensemble le bon rôle pendant la visio "
    "(on a besoin d'un accès en lecture sur les adhérents, adhésions et formations — pas plus)",
)
add_numbered_step(10, "Cliquez Enregistrer puis Créer")

# -- ACTION 3 --
doc.add_heading(
    "Action 3 — M'ouvrir un accès temporaire pour créer et tester les flux", level=2
)

add_paragraph(
    "Les flux qui synchronisent automatiquement Dynamics avec le site (adhérent validé, "
    "mise à jour de profil, désactivation, etc.) seront configurés de mon côté dans votre "
    "environnement Power Automate."
)
add_paragraph(
    "Comme je vais les construire et les tester progressivement, j'ai besoin d'un accès "
    "temporaire me permettant d'avancer de façon autonome, puis de vous restituer une "
    "configuration propre une fois les tests terminés."
)

doc.add_heading("3a — M'ajouter comme utilisateur externe", level=3)
add_numbered_step(1, "Allez sur portal.azure.com")
add_numbered_step(
    2,
    "Dans la barre de recherche, tapez « Utilisateurs » et cliquez sur le résultat sous Microsoft Entra ID",
)
add_numbered_step(3, "Cliquez sur + Inviter un utilisateur externe")
add_numbered_step(4, "Renseignez mon adresse email : (à compléter)")
add_numbered_step(5, "Cliquez sur Inviter")

doc.add_heading(
    "3b — Me donner les droits nécessaires dans l'environnement Power Platform", level=3
)
add_numbered_step(1, "Allez sur admin.powerplatform.microsoft.com")
add_numbered_step(2, "Cliquez sur Environnements")
add_numbered_step(3, "Sélectionnez votre environnement Dynamics")
add_numbered_step(
    4, "Allez dans Paramètres > Utilisateurs + autorisations > Utilisateurs"
)
add_numbered_step(
    5, "Vérifiez que mon compte invité apparaît bien dans la liste"
)
add_numbered_step(6, "Attribuez-moi le rôle Environment Maker")

add_paragraph(
    "Cela me permettra de créer les flux et de faire les tests techniques sans bloquer vos équipes."
)

doc.add_heading("3c — Vérifier la couverture licence", level=3)

add_paragraph(
    "Pour que je puisse créer et tester les flux avec les connecteurs nécessaires "
    "(Dataverse et HTTP), il faut vérifier que mon compte invité pourra bien utiliser "
    "les licences disponibles sur votre tenant."
)

add_numbered_step(1, "Allez sur admin.microsoft.com")
add_numbered_step(2, "Ouvrez Facturation > Licences")
add_numbered_step(3, "Envoyez-nous une capture d'écran de cette page")
add_numbered_step(
    4,
    "Nous vous confirmerons si la couverture est suffisante ou s'il faut prévoir un ajustement",
)

doc.add_heading("3d — Cadre de test", level=3)

add_paragraph(
    "Pendant cette phase, je ferai uniquement des tests ciblés sur quelques enregistrements "
    "de test identifiables."
)

bullets = [
    "Les tests se feront sur un nombre limité de fiches",
    "Les données de test créées dans Dynamics et dans la base du site seront supprimées une fois la recette terminée",
    "Une fois la configuration validée, les flux seront transférés à un compte FNECS pour que vous restiez pleinement autonomes",
]
for b in bullets:
    doc.add_paragraph(b, style="List Bullet")

# -- CE QUI SE PASSE ENSUITE --
doc.add_heading("Ce qui se passe ensuite", level=2)

add_paragraph("Une fois ces 3 actions faites :")

steps_after = [
    "Je configure les flux automatisés dans votre environnement Power Automate — ce sont eux qui préviendront le site quand un adhérent est validé, modifié ou désactivé dans Dynamics",
    "Je teste que tout fonctionne",
    "On fait le transfert : je vous ajoute (vous ou Gérald) comme propriétaire des flux, vous remplacez ma connexion par la vôtre, et je me retire. Comme ça, les flux tournent sous un compte FNECS et ne dépendent plus de mon accès invité",
]
for i, s in enumerate(steps_after, 1):
    add_numbered_step(i, s)

# -- INFORMATIONS A TRANSMETTRE --
doc.add_heading("Informations à nous transmettre", level=2)

add_paragraph(
    "Par un canal sécurisé (message privé Teams, lien de partage sécurisé, etc.) :",
    italic=True,
)

add_simple_table(
    ["Information", "Où la trouver"],
    [
        ["ID d'application (client)", "Page de l'application créée à l'action 1"],
        ["ID de l'annuaire (locataire)", "Page de l'application créée à l'action 1"],
        ["Secret client (la valeur)", "Certificats et secrets de l'action 1"],
        [
            "URL de votre environnement Dynamics",
            "Format : https://fnecs.crm4.dynamics.com ou similaire",
        ],
        ["Capture d'écran des licences", "Action 3c"],
    ],
)

add_paragraph(
    "De notre côté, nous vous fournirons les informations techniques nécessaires à la "
    "configuration des flux (adresses du site, clés d'authentification, format des données attendu)."
)

# -- QUESTION --
doc.add_heading("Une question à trancher ensemble", level=2)

add_paragraph(
    "Quand un adhérent s'inscrit à une formation via le site, est-ce que vous préférez :"
)

options = [
    "A. Que l'inscription remonte automatiquement dans Dynamics (sans action de votre part)",
    "B. Qu'il y ait une étape de validation de votre côté avant que ça apparaisse dans Dynamics",
]
for o in options:
    p = doc.add_paragraph()
    run = p.add_run(o)
    run.bold = True

add_paragraph("On en discutera pendant la visio.")

# -- SIGNATURE --
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("On planifie un créneau de 30 minutes ensemble pour faire les actions 1, 2 et 3 ?").italic = True

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Kevin — La Faabrick Cherdet").bold = True

# -- Sauvegarde --
output_path = r"C:\Users\gaill\La Faabrick Cherdet\La Faabrick Cherdet - Documents\Projets\refonte_site_FNECS\Procedure_Connexion_Dynamics365.docx"
doc.save(output_path)
print(f"Document créé : {output_path}")
